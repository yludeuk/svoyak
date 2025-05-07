import re
import random
from docx import Document
import string
from itertools import combinations
from collections import Counter


def parse_input(filename):
    with open(filename, encoding='utf-8') as f:
        lines = f.read().splitlines()
    return lines


def filter_question_set_to_best_5(questions):
    # Если 5 или меньше — возвращаем как есть
    if len(questions) <= 5:
        return questions
    # Ищем такие 5, что их (отсортированные по возрастанию) можно привести к 10-20-30-40-50
    best = None
    best_score = None
    for comb in combinations(questions, 5):
        sub_prices = [q['price'] for q in comb]
        sorted_prices = sorted(sub_prices)
        # ищем шаг
        diffs = [sorted_prices[i] - sorted_prices[i-1] for i in range(1, 5)]
        cnt = Counter(diffs)
        step, count = cnt.most_common(1)[0]
        # идеал — все шаги одинаковы (тогда count=4)
        # штрафуем за выбросы (максимальный разброс)
        penalty = max(abs(x - step) for x in diffs)
        score = count*10 - penalty
        if best is None or score > best_score:
            best = comb
            best_score = score
    return list(best) if best else questions[:5]


def extract_questions(theme_lines):
    question_re = re.compile(r'^(\d+)[\.\s]')
    service_words = {'http', 'Источник', 'Комментарий', 'Зачет', 'Незачет', 'Ответ'}
    extra_service_words = {
        'Раунд', 'Тема', 'Полуоткрытый', 'Открытый',
        'Закрытый', 'Блок', 'Четвертьфинал', 'Полуфинал', 'Финал',
        'Автор', 'Авторы', 'Редакторы'
    }
    questions = []
    i = 0
    while i < len(theme_lines):
        m = question_re.match(theme_lines[i])
        if m:
            price = int(m.group(1))
            if price % 10 != 0:  # Новый фильтр!
                i += 1
                continue
            rest = theme_lines[i][m.end():].strip()
            words = rest.split()
            first_word = words[0] if words else ""
            second_word = words[1] if len(words) > 1 else ""
            if any(first_word.startswith(sw) for sw in service_words):
                i += 1
                continue
            if (first_word in extra_service_words or
                (second_word and second_word in extra_service_words) or
                (first_word and first_word[0].isupper() and len(words) < 4 and not rest.endswith('?'))):
                i += 1
                continue
            q_start = i
            q_end = i + 1
            while q_end < len(theme_lines):
                m_next = question_re.match(theme_lines[q_end])
                if m_next:
                    rest_next = theme_lines[q_end][m_next.end():].strip()
                    words_next = rest_next.split()
                    first_word_next = words_next[0] if words_next else ""
                    second_word_next = words_next[1] if len(words_next) > 1 else ""
                    if any(first_word_next.startswith(sw) for sw in service_words):
                        q_end += 1
                        continue
                    if (first_word_next in extra_service_words or
                        (second_word_next and second_word_next in extra_service_words) or
                        (first_word_next and first_word_next[0].isupper() and len(words_next) < 4 and not rest_next.endswith('?'))):
                        q_end += 1
                        continue
                    break
                q_end += 1
            questions.append({'price': price, 'start': q_start, 'end': q_end})
            i = q_end
            continue
        i += 1
    questions = filter_question_set_to_best_5(questions)
    return questions

def normalize_prices(prices):
    """
    Сопоставляет реальные номиналы (от 100 и выше) с [10, 20, 30, 40, 50] по возрастанию.
    """
    sorted_prices = sorted(set(prices))
    if len(sorted_prices) > 5:
        sorted_prices = sorted_prices[:5]
    mapping = {old: new for old, new in zip(sorted_prices, [10, 20, 30, 40, 50])}
    return mapping

def extract_answer_block(lines, start_idx, end_idx):
    """
    Возвращает текст вопроса и блока ответа.
    """
    answer_idx = None
    for i in range(start_idx, end_idx):
        if lines[i].startswith('Ответ:'):
            answer_idx = i
            break
    if answer_idx is None:
        return ("", "")
    q_lines = [lines[j] for j in range(start_idx, answer_idx)]
    a_lines = [lines[answer_idx]]
    i = answer_idx+1
    while i < end_idx and not re.match(r'^\d+[\.\s]', lines[i]):
        if lines[i].strip() == '':
            break
        a_lines.append(lines[i])
        i += 1
    q_text = ' '.join(l.strip() for l in q_lines if l.strip())
    a_text = '\n'.join(l for l in a_lines if l.strip())
    return (q_text, a_text)

def process_theme(lines, theme, theme_num):
    """
    Собирает одну тему в финальный текстовый блок.
    """
    theme_lines = lines[theme['start']+1:theme['end']]
    questions = extract_questions(theme_lines)
    if len(questions) == 0:
        return ""
    prices = [q['price'] for q in questions]
    mapping = normalize_prices(prices)
    result = []
    for idx, q in enumerate(questions):
        price_norm = mapping[q['price']]
        q_text, a_text = extract_answer_block(theme_lines, q['start'], q['end'])
        q_text = re.sub(r'^\d+[\.\s]', '', q_text).strip()
        if not q_text and not a_text:
            continue
        result.append(f"{price_norm}. {q_text}\n\n{a_text}\n")
    out = [f"{theme_num}. {theme['name']}\n"]
    out += [item for item in result]
    return '\n'.join(out).strip()


def extract_themes(lines):
    themes = []
    i = 0
    theme_re1 = re.compile(r'Тема\s*(\d+)\.?\s*(.+)')
    theme_re2 = re.compile(r'(\d+)\s*Тема: (.+)')
    while i < len(lines):
        m1 = theme_re1.match(lines[i])
        m2 = theme_re2.match(lines[i])
        if m1:
            num = int(m1.group(1))
            name = m1.group(2).strip()
            themes.append({'num': num, 'name': name, 'start': i})
        elif m2:
            num = int(m2.group(1))
            name = m2.group(2).strip()
            themes.append({'num': num, 'name': name, 'start': i})
        i += 1
    for idx in range(len(themes)):
        themes[idx]['end'] = themes[idx + 1]['start'] if idx+1 < len(themes) else len(lines)
    return themes


import re

def is_question_line(line):
    """
    Вопросная строка: начинается с числа (от 10 до 10000, кратного 10) и точки/пробела.
    """
    m = re.match(r'^(\d+)[\.\s]', line.strip())
    if not m:
        return False
    price = int(m.group(1))
    # Только номиналы от 10 до 10000 и кратные 10
    if price % 10 != 0 or price < 10 or price > 10000:
        return False
    return True

def find_next_questions_block(lines, start, block_size=5, max_window=100):
    """
    Ищет блок из block_size вопросов (строки, начинающиеся на номинал),
    возвращает списком индексы их в пределах окна max_window строк.
    """
    indices = []
    for i in range(start, min(len(lines), start + max_window)):
        if is_question_line(lines[i]):
            indices.append(i)
            if len(indices) == block_size:
                return indices
    return []


def clean_theme_name(line):
    """
    Убирает варианты 'Тема', номер, точки, двоеточия и возвращает только чистое название темы.
    """
    # регулярка: опциональный номер, слово Тема, опциональный номер, точки, двоеточия, пробелы, потом всё остальное
    m = re.match(r'^(?:\d+\s*)?(?:Тема)?\s*(?:\d+)?[\.:\s-]*([^\d\W][\w\W]*)$', line.strip(), re.IGNORECASE)
    if m:
        name = m.group(1).strip()
    else:
        name = line.strip()
    return name


def flexible_extract_themes(lines):
    themes = []
    i = 0
    num_theme = 1
    while i < len(lines):
        block = find_next_questions_block(lines, i, block_size=5, max_window=30)
        if block:
            t_idx = block[0] - 1
            # Если строка 'Автор', идём выше
            while t_idx >= 0 and (not lines[t_idx].strip() or lines[t_idx].strip().lower().startswith('автор')):
                t_idx -= 1
            # Если после пропуска авторов всё равно оказались на пустой строке — идём выше
            while t_idx >= 0 and not lines[t_idx].strip():
                t_idx -= 1
            theme_name = clean_theme_name(lines[t_idx]) if t_idx >= 0 else f'Тема {num_theme}'
            themes.append({'num': num_theme, 'name': theme_name, 'start': t_idx, 'q_start': block[0]})
            num_theme += 1
            i = block[-1] + 1
        else:
            i += 1
    for idx in range(len(themes)):
        themes[idx]['end'] = themes[idx + 1]['start'] if idx + 1 < len(themes) else len(lines)
    return themes


def auto_split(total, min_size=9, max_size=12):
    result = []
    rest = total
    while rest > 0:
        if total > 100:
            # После 100 тем только 10 или 11
            size = 11 if rest % 11 == 0 or rest == 11 else 10
            if rest < 10:
                size = rest
        else:
            # До 100 тем
            # Не даём блоков <9 и >12
            if rest <= max_size:
                size = rest
            elif rest % max_size == 0 or rest % (max_size-1) == 0:
                size = max_size
            elif rest % min_size == 0 or rest % (min_size+1) == 0:
                size = min_size
            elif rest >= max_size*2:
                size = max_size
            else:
                size = min(max_size, rest)
        result.append(size)
        rest -= size
    return result


def check_user_split(split, total, min_size=9, max_size=12):
    if sum(split) != total:
        return False
    for s in split:
        if s < min_size or s > max_size:
            return False
    return True


def split_themes(themes, user_split=None, min_size=9, max_size=12):
    total = len(themes)
    if user_split and check_user_split(user_split, total, min_size, max_size):
        sizes = user_split
    else:
        sizes = even_split(total, min_size, max_size)
    blocks = []
    idx = 0
    for sz in sizes:
        blocks.append(themes[idx:idx+sz])
        idx += sz
    return blocks


def even_split(total, min_size=9, max_size=12):
    for blocks in range(total // max_size, total // min_size + 2):
        if blocks == 0:
            continue
        base = total // blocks
        plus = total % blocks
        sizes = [base+1]*plus + [base]*(blocks-plus)
        if all(min_size <= s <= max_size for s in sizes):
            return sorted(sizes)
    sizes = [max_size] * (total // max_size)
    if total % max_size:
        sizes.append(total % max_size)
    return sorted(sizes)


def random_prefix(length=4):
    return ''.join(random.choices(string.ascii_lowercase, k=length))


def save_blocks_to_docx(block_texts, prefix=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if prefix is None:
        prefix = random_prefix()
    for i, text in enumerate(block_texts, 1):
        doc = Document()
        for line in text.split('\n'):
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run('')
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            # Для корректного отображения кириллицы в некоторых Word'ах
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            # Минимальные интервалы
            p_format = p.paragraph_format
            p_format.space_after = Pt(0)
            p_format.space_before = Pt(0)
            p_format.line_spacing = 1.0  # Одинарный интервал
        doc.save(f"{prefix}_{i}.docx")



def main():
    lines = parse_input('input.txt')
    themes = flexible_extract_themes(lines)

    theme_blocks = split_themes(themes, [10, 10, 10, 10, 10, 10, 10, 10])
    theme_counter = 1
    block_texts = []
    for block in theme_blocks:
        outs = []
        for theme in block:
            block_text = process_theme(lines, theme, theme_counter)
            outs.append(block_text)
            theme_counter += 1
        block_texts.append('\n\n'.join(outs).strip())
    # Сохраняем в docx
    save_blocks_to_docx(block_texts)

    # outputs = []
    # theme_counter = 1
    # for theme in themes:
    #     block = process_theme(lines, theme, theme_counter)
    #     outputs.append(block)
    #     theme_counter += 1
    # txt = '\n\n'.join(outputs).strip() + '\n'
    # with open('parsed_output.txt', 'w', encoding='utf-8') as f:
    #     f.write(txt)


if __name__ == '__main__':
    main()
